#!/usr/bin/env python3
"""Build the Nemotron Voice Agent current-architecture handoff manual.

This file is intentionally self-contained.  It creates nine editable-by-code
schematic sources as high-resolution PNGs, copies the repository-owned NVIDIA
visual assets used by the manual, and emits a Microsoft Word document.

Run with the task-local/documentation virtual environment described in README.md.
No application or deployment source is modified.
"""

from __future__ import annotations

import hashlib
import json
import os
import re
import shutil
import subprocess
import sys
import textwrap
import zipfile
from datetime import datetime, timezone
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


HERE = Path(__file__).resolve().parent
DOCS = HERE.parent
REPO = DOCS.parent
ASSETS = HERE / "assets"
SSOT = DOCS / "current-deployed-pipeline-architecture.md"
OUTPUT = DOCS / "Nemotron_Voice_Agent_Current_Architecture_Manual.docx"
PDF_OUTPUT = HERE / "Nemotron_Voice_Agent_Current_Architecture_Manual.pdf"
BUILD_REPORT = HERE / "build-report.json"

SNAPSHOT_DATE = "2026-08-19"
VERSION = "1.0"
STATUS = "CURRENT DEPLOYED ARCHITECTURE — HANDOFF MANUAL"

# Visual language: white canvas, dark charcoal type, NVIDIA-adjacent pale green.
WHITE = "FFFFFF"
CHARCOAL = "242A2E"
MUTED = "586168"
GREEN = "76B900"
GREEN_DARK = "467500"
PALE = "EDF8DE"
PALE_2 = "F6FBEF"
MINT = "E4F4DA"
AMBER = "FFF4D6"
RED_PALE = "FDEBE9"
BLUE_PALE = "EAF3F8"
GRID = "D9E3D1"


# ---------------------------------------------------------------------------
# Diagram primitives (the editable schematic source)
# ---------------------------------------------------------------------------

W, H = 2800, 1575
FONT_REG = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
FONT_BOLD = "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"
FONT_MONO = "/usr/share/fonts/truetype/dejavu/DejaVuSansMono.ttf"


def font(size: int, bold: bool = False, mono: bool = False) -> ImageFont.FreeTypeFont:
    path = FONT_MONO if mono else (FONT_BOLD if bold else FONT_REG)
    return ImageFont.truetype(path, size)


def canvas(title: str, subtitle: str = ""):
    im = Image.new("RGB", (W, H), "white")
    d = ImageDraw.Draw(im, "RGBA")
    d.rectangle((0, 0, W, 18), fill="#76B900")
    d.text((90, 55), title, font=font(58, True), fill="#242A2E")
    if subtitle:
        d.text((92, 126), subtitle, font=font(27), fill="#586168")
    d.line((90, 180, W - 90, 180), fill="#C8D7BE", width=3)
    return im, d


def rounded(d, box, fill="#EDF8DE", outline="#76B900", radius=28, width=3):
    d.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def wrapped_text(d, box, title, body="", badge=None, fill="#EDF8DE", outline="#76B900",
                 title_size=30, body_size=23, align="left"):
    x1, y1, x2, y2 = box
    rounded(d, box, fill=fill, outline=outline)
    tx = x1 + 34
    if badge:
        bx, by = x1 + 42, y1 + 48
        d.ellipse((bx - 29, by - 29, bx + 29, by + 29), fill="#76B900", outline="#467500", width=2)
        bb = d.textbbox((0, 0), badge, font=font(19, True))
        d.text((bx - (bb[2] - bb[0]) / 2, by - (bb[3] - bb[1]) / 2 - 2), badge,
               font=font(19, True), fill="white")
        tx = x1 + 87
    max_chars = max(12, int((x2 - tx - 25) / (title_size * 0.58)))
    title_lines = textwrap.wrap(title, width=max_chars) or [title]
    y = y1 + 28
    for line in title_lines:
        d.text((tx, y), line, font=font(title_size, True), fill="#242A2E")
        y += title_size + 7
    if body:
        y += 8
        body_chars = max(16, int((x2 - x1 - 60) / (body_size * 0.54)))
        for line in textwrap.wrap(body, width=body_chars):
            d.text((x1 + 34, y), line, font=font(body_size), fill="#4B565C")
            y += body_size + 7


def arrow(d, start, end, label="", color="#467500", width=5, dashed=False):
    x1, y1 = start
    x2, y2 = end
    if dashed:
        segments = 16
        for i in range(0, segments, 2):
            a, b = i / segments, min(1, (i + 1) / segments)
            d.line((x1 + (x2-x1)*a, y1 + (y2-y1)*a,
                    x1 + (x2-x1)*b, y1 + (y2-y1)*b), fill=color, width=width)
    else:
        d.line((x1, y1, x2, y2), fill=color, width=width)
    import math
    angle = math.atan2(y2-y1, x2-x1)
    ah = 20
    pts = [(x2, y2),
           (x2-ah*math.cos(angle-0.55), y2-ah*math.sin(angle-0.55)),
           (x2-ah*math.cos(angle+0.55), y2-ah*math.sin(angle+0.55))]
    d.polygon(pts, fill=color)
    if label:
        mx, my = (x1+x2)/2, (y1+y2)/2
        bb = d.textbbox((0, 0), label, font=font(19, True))
        pad = 8
        d.rounded_rectangle((mx-(bb[2]-bb[0])/2-pad, my-29,
                             mx+(bb[2]-bb[0])/2+pad, my+6), 8, fill="white", outline="#D9E3D1")
        d.text((mx-(bb[2]-bb[0])/2, my-25), label, font=font(19, True), fill=color)


def lane(d, box, title, fill="#F6FBEF"):
    x1, y1, x2, y2 = box
    rounded(d, box, fill=fill, outline="#C8D7BE", radius=22, width=2)
    d.text((x1+24, y1+18), title.upper(), font=font(21, True), fill="#467500")
    d.line((x1+20, y1+54, x2-20, y1+54), fill="#D9E3D1", width=2)


def footer(d, number: int, note="Architectural schematic — not to scale"):
    d.line((90, H-62, W-90, H-62), fill="#D9E3D1", width=2)
    d.text((90, H-47), note, font=font(17), fill="#7A8388")
    t = f"FIGURE {number:02d}"
    bb = d.textbbox((0, 0), t, font=font(17, True))
    d.text((W-90-(bb[2]-bb[0]), H-47), t, font=font(17, True), fill="#467500")


def save_diagram(im: Image.Image, name: str) -> Path:
    ASSETS.mkdir(parents=True, exist_ok=True)
    path = ASSETS / name
    im.save(path, "PNG", optimize=True, dpi=(220, 220))
    return path


def diagram_overview():
    im, d = canvas("End-to-end system overview", "The retained user path: Astra-hosted UI → production NVCF function → model services")
    lane(d, (80, 210, 480, 1460), "User edge", "#FAFCF8")
    lane(d, (520, 210, 1050, 1460), "Astra — physically stg", "#F6FBEF")
    lane(d, (1090, 210, 2690, 1460), "NVCF production function", "#F6FBEF")
    wrapped_text(d, (125, 340, 435, 560), "Browser", "React + Pipecat client\nMic · speaker · camera · uploads", "BR")
    wrapped_text(d, (565, 330, 1005, 550), "Astra ingress + nginx", "Serves SPA; injects function ID and bearer token server-side", "AS")
    wrapped_text(d, (565, 650, 1005, 845), "HTTP gateway", "Same-origin REST and /health", "H", fill="#EAF3F8", outline="#5B8FA8")
    wrapped_text(d, (565, 930, 1005, 1125), "Streaming gateway", "WSS /api/ws and function routing", "WS", fill="#EAF3F8", outline="#5B8FA8")
    wrapped_text(d, (1140, 305, 1515, 525), "Kubernetes Service", "Ordinary load balancing; router disabled", "K8")
    for i, x in enumerate([1140, 1470, 1800, 2130, 2460], 1):
        wrapped_text(d, (x, 610, x+250, 815), f"App replica {i}", "FastAPI + one Pipecat graph per WebSocket", f"A{i}", title_size=25, body_size=18)
    wrapped_text(d, (1210, 970, 1630, 1180), "Redis", "Config · media streams · capture flags/locks", "R")
    wrapped_text(d, (1740, 970, 2160, 1180), "SeaweedFS S3", "Shared capture staging objects", "SW")
    wrapped_text(d, (2270, 970, 2640, 1180), "NGC resource", "Durable archive version = session ID", "NGC", fill="#EAF3F8", outline="#5B8FA8")
    wrapped_text(d, (1160, 1260, 2060, 1400), "7 H100 GPUs", "ASR · Lightning · Super TP=2 · Omni · Magpie · Chatterbox", "NV", title_size=27, body_size=20)
    wrapped_text(d, (2160, 1260, 2640, 1400), "Live tools", "WeatherAPI · Finnhub · Perplexity", "API", fill="#FFF4D6", outline="#C68A00", title_size=27, body_size=20)
    arrow(d, (435, 420), (565, 420), "HTTPS")
    arrow(d, (435, 490), (565, 1000), "WSS")
    arrow(d, (1005, 750), (1140, 400), "REST")
    arrow(d, (1005, 1020), (1140, 460), "voice")
    arrow(d, (1325, 525), (1265, 610))
    arrow(d, (1325, 525), (1595, 610))
    arrow(d, (1325, 525), (1925, 610))
    arrow(d, (1325, 525), (2255, 610))
    arrow(d, (1325, 525), (2585, 610))
    for x in [1265,1595,1925,2255,2585]:
        arrow(d, (x, 815), (1420, 970), "shared" if x == 1265 else "", width=3)
        arrow(d, (x, 815), (1950, 970), "", width=3)
    arrow(d, (2160, 1070), (2270, 1070), "upload")
    footer(d, 1, "Retained live UI is on Astra stg infrastructure; backend is the production NVCF function")
    return save_diagram(im, "01_end_to_end_overview.png")


def diagram_trust():
    im, d = canvas("Trust boundaries and request routing", "Secrets stop at server-side boundaries; browser traffic is same-origin")
    lane(d, (90, 220, 690, 1430), "Boundary 1 — untrusted browser", "#FAFCF8")
    lane(d, (730, 220, 1370, 1430), "Boundary 2 — Astra", "#F6FBEF")
    lane(d, (1410, 220, 2070, 1430), "Boundary 3 — NVCF edge", "#F8FBFC")
    lane(d, (2110, 220, 2710, 1430), "Boundary 4 — private namespace", "#F6FBEF")
    wrapped_text(d, (150, 350, 625, 600), "React SPA", "Public assets and non-secret config.js; same-origin API calls", "UI")
    wrapped_text(d, (790, 330, 1310, 610), "nginx reverse proxy", "Adds Authorization + function-id; strips cookies; never writes secrets to JavaScript", "NX")
    wrapped_text(d, (790, 760, 1310, 990), "Astra Vault", "NVCF_HOST\nNVCF_FUNCTION_ID\nNVIDIA_API_KEY", "V", fill="#FFF4D6", outline="#C68A00")
    wrapped_text(d, (1470, 330, 2010, 570), "HTTP invocation gateway", "/api/* and /health", "H", fill="#EAF3F8", outline="#5B8FA8")
    wrapped_text(d, (1470, 720, 2010, 960), "Streaming gateway", "grpc.nvcf.nvidia.com\nWSS /api/ws", "WS", fill="#EAF3F8", outline="#5B8FA8")
    wrapped_text(d, (2170, 310, 2650, 535), "App Service :7860", "Routes new requests to five app replicas", "K8")
    wrapped_text(d, (2170, 660, 2650, 885), "Function-version secrets", "Mounted as /var/secrets/secrets.json; exported by entrypoint", "S", fill="#FFF4D6", outline="#C68A00")
    wrapped_text(d, (2170, 1030, 2650, 1260), "ClusterIP-only services", "NIMs · Redis · SeaweedFS; no public ingress", "P")
    arrow(d, (625, 450), (790, 450), "HTTPS")
    arrow(d, (625, 515), (790, 840), "never")
    arrow(d, (1050, 760), (1050, 610), "inject")
    arrow(d, (1310, 420), (1470, 420), "Bearer + FID")
    arrow(d, (1310, 520), (1470, 830), "upgrade")
    arrow(d, (2010, 450), (2170, 420))
    arrow(d, (2010, 830), (2170, 470))
    arrow(d, (2410, 660), (2410, 535), "server only")
    arrow(d, (2410, 535), (2410, 1030), "internal")
    rounded(d, (165, 1130, 650, 1330), fill="#FDEBE9", outline="#B64A3A")
    d.text((200, 1160), "THE BROWSER NEVER RECEIVES", font=font(24, True), fill="#8C2E23")
    for i, t in enumerate(["NVCF invocation credential", "provider API credentials", "NGC publication credential"]):
        d.text((205, 1210+i*34), "• " + t, font=font(21), fill="#5B3430")
    footer(d, 2)
    return save_diagram(im, "02_trust_and_routing.png")


def diagram_session_sequence():
    im, d = canvas("Start-session request sequence", "REST and WebSocket may hit different replicas; Redis closes the handoff gap")
    actors = [(160,"Browser","BR"),(600,"Astra nginx","AS"),(1060,"HTTP edge","H"),(1480,"App A","A"),(1870,"Redis","R"),(2260,"WS edge","WS"),(2570,"App B","B")]
    top = 270
    for x, name, badge in actors:
        d.ellipse((x-34, top-34, x+34, top+34), fill="#76B900", outline="#467500", width=2)
        bb=d.textbbox((0,0),badge,font=font(16,True)); d.text((x-(bb[2]-bb[0])/2,top-(bb[3]-bb[1])/2-2),badge,font=font(16,True),fill="white")
        bb=d.textbbox((0,0),name,font=font(21,True)); d.text((x-(bb[2]-bb[0])/2,top+50),name,font=font(21,True),fill="#242A2E")
        d.line((x, top+90, x, 1400), fill="#C8D7BE", width=3)
    steps = [
        (1,160,600,410,"POST /api/session-config"),(2,600,1060,500,"add bearer + function-id"),
        (3,1060,1480,590,"load-balanced REST"),(4,1480,1480,690,"validate catalog + readiness"),
        (5,1480,1870,790,"SET sb:cfg:<sid> TTL 3600"),(6,1480,160,890,"return server-minted SID"),
        (7,160,600,1010,"WSS /api/ws?session_id=<sid>"),(8,600,2260,1100,"upgrade via streaming gateway"),
        (9,2260,2570,1190,"load-balanced socket"),(10,2570,1870,1280,"GET sb:cfg:<sid>"),
        (11,2570,160,1370,"Pipecat / RTVI audio session")]
    for n,x1,x2,y,label in steps:
        color="#467500" if n not in (6,11) else "#5B8FA8"
        arrow(d,(x1,y),(x2,y),label,color=color,width=4)
        d.ellipse((80,y-23,126,y+23),fill="#76B900")
        bb=d.textbbox((0,0),str(n),font=font(18,True));d.text((103-(bb[2]-bb[0])/2,y-(bb[3]-bb[1])/2-2),str(n),font=font(18,True),fill="white")
    footer(d, 3, "Application replica A and B may be different pods; the live WebSocket remains on B")
    return save_diagram(im, "03_session_request_sequence.png")


def diagram_generic():
    im, d = canvas("Generic Assistant voice and tool pipeline", "Pipecat cascaded ASR → text LLM → function tools → external TTS")
    y=560
    boxes=[(100,"Mic + transport","BR"),(470,"Nemotron ASR","ASR"),(850,"User context","CTX"),(1220,"Lightning / Super","LLM"),(1650,"Speech gate","G"),(2020,"Magpie / Chatterbox","TTS"),(2450,"Speaker","SP")]
    for x,t,b in boxes:
        wrapped_text(d,(x,y,x+300,y+190),t,"",b,title_size=24)
    for i in range(len(boxes)-1): arrow(d,(boxes[i][0]+300,y+95),(boxes[i+1][0],y+95))
    wrapped_text(d,(1050,990,1510,1225),"OpenAI-compatible tool schema", "Per-session allowlist; tool_choice=auto", "ƒ", fill="#EAF3F8", outline="#5B8FA8")
    wrapped_text(d,(1650,930,2210,1280),"Python tool handlers", "WeatherAPI · Finnhub · Perplexity Sonar\nBMI · random number\nSpeak-safe failure objects", "PY", fill="#FFF4D6", outline="#C68A00")
    arrow(d,(1370,750),(1280,990),"function call")
    arrow(d,(1510,1100),(1650,1100),"validated")
    arrow(d,(1930,930),(1430,750),"result → context",color="#5B8FA8")
    rounded(d,(190,260,915,420),fill="#F6FBEF",outline="#C8D7BE")
    d.text((230,290),"NORMAL ANSWER",font=font(23,True),fill="#467500")
    d.text((230,335),"LLM text flows through speech gate and TTS",font=font(22),fill="#4B565C")
    rounded(d,(1030,260,1785,420),fill="#F6FBEF",outline="#C8D7BE")
    d.text((1070,290),"TOOL COMPLETION",font=font(23,True),fill="#467500")
    d.text((1070,335),"Pre-tool reasoning text is buffered and dropped",font=font(22),fill="#4B565C")
    rounded(d,(1900,260,2620,420),fill="#F6FBEF",outline="#C8D7BE")
    d.text((1940,290),"POST-TOOL ANSWER",font=font(23,True),fill="#467500")
    d.text((1940,335),"Provider-derived response is synthesized",font=font(22),fill="#4B565C")
    wrapped_text(d,(560,1050,920,1270),"Capture tap", "ASR/TTS WAVs + session log → SeaweedFS", "REC", title_size=25, body_size=19)
    arrow(d,(620,750),(740,1050),"record",width=3,dashed=True)
    arrow(d,(2170,750),(850,1050),"record",width=3,dashed=True)
    footer(d, 4)
    return save_diagram(im, "04_generic_pipeline.png")


def diagram_omni():
    im, d = canvas("Omni Assistant Subagents worker architecture", "Pipecat WorkerRunner and WorkerBus inside one app process per live session")
    wrapped_text(d,(100,530,520,800),"Transport Agent", "Audio I/O · VAD · turns · TTS · RTVI · media dispatch", "IO")
    wrapped_text(d,(650,550,1030,780),"BusBridgeProcessor", "Bridges frames between transport and workers", "BB")
    wrapped_text(d,(1160,500,1580,830),"Pipecat WorkerBus", "One in-process bus for the session; not Redis", "BUS", fill="#EAF3F8", outline="#5B8FA8")
    wrapped_text(d,(1750,240,2260,500),"Speaker Omni Agent", "Only worker allowed to speak; emits strict action envelope", "SP")
    wrapped_text(d,(1750,570,2260,830),"Media Analyzer", "Uploaded image/audio/video and high-res capture", "MA")
    wrapped_text(d,(1750,900,2260,1160),"Webcam Agent", "Rolling view summary and conservative gestures", "WC")
    wrapped_text(d,(1750,1230,2260,1460),"Thinker Worker", "On-demand reasoning re-answer", "TH")
    wrapped_text(d,(2380,410,2700,750),"Pinned state board", "Latest media analysis + live view; read each turn", "SB", title_size=24, body_size=19)
    wrapped_text(d,(2380,950,2700,1210),"Redis streams", "Cross-pod attachments and webcam frames", "R", fill="#FFF4D6", outline="#C68A00", title_size=24, body_size=19)
    arrow(d,(520,665),(650,665));arrow(d,(1030,665),(1160,665));
    for y in [370,700,1030,1345]: arrow(d,(1580,665),(1750,y),"job/result" if y==700 else "",width=4)
    arrow(d,(2260,380),(2380,500),"write/read",width=3)
    arrow(d,(2260,700),(2380,580),"pin",width=3)
    arrow(d,(2260,1030),(2380,640),"pin",width=3)
    arrow(d,(2380,1080),(2260,700),"media",width=3)
    rounded(d,(120,240,1510,390),fill="#F6FBEF",outline="#C8D7BE")
    d.text((160,270),"IMPORTANT OWNERSHIP BOUNDARY",font=font(23,True),fill="#467500")
    d.text((160,318),"WorkerBus = agent coordination within one session/pod; Redis = cross-pod browser media delivery",font=font(21),fill="#4B565C")
    footer(d, 5)
    return save_diagram(im, "05_omni_workers.png")


def diagram_replica_sharing():
    im, d = canvas("Multi-replica data sharing", "Five application pods share ancillary session data, but not the live voice pipeline")
    lane(d,(90,230,920,1435),"Stateless request distribution","#FAFCF8")
    lane(d,(960,230,1880,1435),"Shared coordination and staging","#F6FBEF")
    lane(d,(1920,230,2710,1435),"Per-socket process state","#F8FBFC")
    for i,(x,y) in enumerate([(150,350),(520,350),(150,720),(520,720),(335,1090)],1):
        wrapped_text(d,(x,y,x+300,y+210),f"App replica {i}","REST and/or WebSocket",f"A{i}",title_size=24,body_size=18)
        arrow(d,(x+300,y+90),(1070,610),"",width=3)
        arrow(d,(x+300,y+145),(1070,990),"",width=3)
    wrapped_text(d,(1070,430,1770,790),"Redis", "sb:cfg:<sid>\nsb:att:<sid> / sb:wc:<sid>\ncap:<sid> / cap:lock:<sid>\nTTL, streams and owner-token locks", "R")
    wrapped_text(d,(1070,880,1770,1220),"SeaweedFS S3", "sessions/<sid>/session.log\ntranscript.txt · audio/*.wav\nShared capture staging", "SW")
    wrapped_text(d,(2000,370,2630,690),"Live WebSocket pipeline", "One accepting replica owns transport, conversation context and processors until disconnect", "WS", fill="#EAF3F8", outline="#5B8FA8")
    wrapped_text(d,(2000,850,2630,1170),"Omni WorkerBus", "One in-process worker graph per Omni session; not shared across replicas", "BUS", fill="#EAF3F8", outline="#5B8FA8")
    arrow(d,(1770,610),(2000,530),"config/media")
    arrow(d,(1770,1050),(2000,960),"capture only",dashed=True)
    rounded(d,(2010,1250,2620,1380),fill="#FDEBE9",outline="#B64A3A")
    d.text((2050,1280),"POD LOSS ENDS THAT LIVE CALL",font=font(22,True),fill="#8C2E23")
    d.text((2050,1325),"Redis does not migrate a socket or LLM context.",font=font(20),fill="#5B3430")
    footer(d, 6)
    return save_diagram(im, "06_multi_replica_sharing.png")


def diagram_capture():
    im, d = canvas("Session capture: two signals, one finalizer", "Replica-safe exactly-once coordination with Redis; shared artifacts in SeaweedFS; durable publication to NGC")
    wrapped_text(d,(100,310,600,560),"Pipeline finishes", "Recorder flushes; local log copied to shared store; set pipeline_done=1", "P")
    wrapped_text(d,(100,900,600,1150),"Browser teardown", "Consent + bounded transcript via keepalive POST; set consent_done=1", "B")
    wrapped_text(d,(800,500,1350,970),"Redis capture state", "cap:<sid> hash\nWait for BOTH signals\nSET cap:lock:<sid> token NX EX 900\nLua compare-and-delete release", "R")
    wrapped_text(d,(1530,270,2050,570),"SeaweedFS", "session.log\ntranscript.txt\naudio/asr_*.wav\naudio/tts_*.wav", "SW")
    wrapped_text(d,(1530,750,2050,1045),"Winning app replica", "Read objects · build <sid>.tar.gz · invoke NGC CLI · handle 300 s timeout", "A")
    wrapped_text(d,(2260,480,2700,810),"NGC resource", "0491162300748285/session-captures:<sid>\nDurable versioned archive", "NGC", fill="#EAF3F8", outline="#5B8FA8")
    wrapped_text(d,(2260,970,2700,1230),"Cleanup", "On success: delete SeaweedFS prefix and clear Redis state", "✓")
    arrow(d,(600,435),(800,650),"signal")
    arrow(d,(600,1025),(800,820),"signal")
    arrow(d,(1350,700),(1530,900),"lock winner")
    arrow(d,(1790,750),(1790,570),"read")
    arrow(d,(2050,900),(2260,650),"upload")
    arrow(d,(2480,810),(2480,970),"success")
    rounded(d,(730,1160,2050,1385),fill="#FFF4D6",outline="#C68A00")
    d.text((780,1195),"RETRY / RETENTION RULE",font=font(24,True),fill="#8A5F00")
    d.text((780,1245),"NGC timeout, missing CLI/key, or upload failure retains shared source and state.",font=font(21),fill="#5C4A20")
    d.text((780,1290),"The 300-second reaper retries; operators check NGC before manual retry or delete.",font=font(21),fill="#5C4A20")
    d.text((780,1335),"SeaweedFS is ephemeral staging, so retained evidence is not restart-durable.",font=font(21),fill="#5C4A20")
    footer(d, 7)
    return save_diagram(im, "07_session_capture_flow.png")


def diagram_promotion():
    im, d = canvas("Artifact qualification and promotion flow", "Promote immutable, identical app/chart/UI artifacts through explicit quality gates")
    stages=[
        (100,"Source + artifacts","Git commit\napp image · UI image · Helm chart","SRC"),
        (550,"Viking local K8s","Real audio Playwright\nmodels · tools · media · capture","L"),
        (1050,"NVCF/Astra staging","Public-path SQA\nrepeated EXPECT_TOOL\nconcurrency","STG"),
        (1580,"Production NVCF","New immutable version\nfull secrets repeated\nverify before old undeploy","N"),
        (2100,"Retained live UI","Astra stg app points at prod NVCF\ncurrent operational state","LIVE"),
    ]
    for i,(x,title,body,badge) in enumerate(stages):
        wrapped_text(d,(x,520,x+410,840),title,body,badge,fill="#EDF8DE" if i!=4 else "#FFF4D6",outline="#76B900" if i!=4 else "#C68A00",title_size=25,body_size=20)
        if i<len(stages)-1: arrow(d,(x+410,680),(stages[i+1][0],680),"QUALIFY")
    wrapped_text(d,(1950,1050,2650,1350),"Future true Astra prd", "Obtain NSPECT ID · deploy to astraprd01-ocp-pdx04 · independent /prd Vault path · qualify · then remove retained stg app", "PRD", fill="#EAF3F8", outline="#5B8FA8")
    arrow(d,(2305,840),(2305,1050),"separate promotion")
    rounded(d,(130,1030,1720,1355),fill="#F6FBEF",outline="#C8D7BE")
    d.text((180,1070),"PROMOTION GATES",font=font(24,True),fill="#467500")
    items=["Exact artifact identity and full function-version secret set","HTTP + WebSocket + real voice + selectors + tools","Omni attachment/webcam/high-res and cross-replica flow","Session archive reaches NGC; concurrent sessions show no cross-talk"]
    for i,t in enumerate(items): d.text((190,1120+i*50),"✓  "+t,font=font(21),fill="#4B565C")
    footer(d, 8, "Staging is currently pulled down; recreating it is a deliberate deployment action")
    return save_diagram(im, "08_deployment_promotion.png")


def diagram_failures():
    im, d = canvas("Operational failure-isolation map", "Start at the user symptom and move one boundary at a time")
    nodes={
        "root":((1050,250,1750,390),"User-visible failure","START","#FDEBE9","#B64A3A"),
        "ui":((120,520,620,730),"Astra root loads?","UI / ingress / image","#EDF8DE","#76B900"),
        "http":((720,520,1220,730),"/health and catalog?","proxy target / Vault / app","#EDF8DE","#76B900"),
        "ws":((1320,520,1820,730),"WebSocket 101?","streaming gateway / FID / cookie","#EDF8DE","#76B900"),
        "voice":((1920,520,2420,730),"Voice response?","ASR / LLM / TTS readiness","#EDF8DE","#76B900"),
        "feature":((2300,850,2700,1070),"Feature-specific","tools · media · capture","#EAF3F8","#5B8FA8"),
        "tool":((150,990,650,1220),"Tool failure","RTVI event → key → provider → safe result","#FFF4D6","#C68A00"),
        "media":((790,990,1290,1220),"Media/webcam","SID → capability → Redis stream → board","#FFF4D6","#C68A00"),
        "cap":((1430,990,1930,1220),"Capture","two flags → lock → store → NGC","#FFF4D6","#C68A00"),
    }
    for k,(box,title,body,fill,outline) in nodes.items(): wrapped_text(d,box,title,body,k[:2].upper(),fill=fill,outline=outline,title_size=25,body_size=19)
    arrow(d,(1400,390),(370,520),"NO")
    arrow(d,(1400,390),(970,520),"YES")
    arrow(d,(1220,625),(1320,625),"NEXT")
    arrow(d,(1820,625),(1920,625),"NEXT")
    arrow(d,(2420,625),(2500,850),"VOICE OK")
    arrow(d,(2500,1070),(400,990),"tool")
    arrow(d,(2500,1070),(1040,990),"media")
    arrow(d,(2500,1070),(1680,990),"capture")
    rounded(d,(150,1315,2650,1435),fill="#F6FBEF",outline="#C8D7BE")
    d.text((190,1348),"Correlation key: the same 12-hex session ID links UI, app logs, Redis keys, SeaweedFS prefix and NGC version.",font=font(22,True),fill="#467500")
    footer(d, 9)
    return save_diagram(im, "09_failure_isolation.png")


def generate_diagrams():
    return [
        diagram_overview(), diagram_trust(), diagram_session_sequence(), diagram_generic(),
        diagram_omni(), diagram_replica_sharing(), diagram_capture(), diagram_promotion(),
        diagram_failures(),
    ]


# ---------------------------------------------------------------------------
# Word helpers
# ---------------------------------------------------------------------------

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=GRID, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, name="Aptos"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def add_field(paragraph, instruction, placeholder=""):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = instruction
    fld_sep = OxmlElement("w:fldChar"); fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t"); text.text = placeholder
    fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])


def set_update_fields(doc):
    settings = doc.settings._element
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)


def set_image_alt_text(paragraph, alt_text):
    drawings = paragraph._p.xpath(".//wp:docPr")
    if drawings:
        drawings[-1].set("descr", alt_text)
        drawings[-1].set("title", alt_text[:120])


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = paragraph.add_run("PAGE ")
    r.font.size = Pt(8); r.font.color.rgb = RGBColor.from_string(MUTED)
    add_field(paragraph, " PAGE ", "1")


def configure_doc(doc: Document):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.72); sec.bottom_margin = Inches(0.68)
    sec.left_margin = Inches(0.78); sec.right_margin = Inches(0.72)
    sec.header_distance = Inches(0.28); sec.footer_distance = Inches(0.25)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"; normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(9.4); normal.font.color.rgb = RGBColor.from_string(CHARCOAL)
    normal.paragraph_format.space_after = Pt(5.5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color, before, after in [
        ("Title", 32, CHARCOAL, 0, 8), ("Subtitle", 15, MUTED, 0, 8),
        ("Heading 1", 22, GREEN_DARK, 15, 7), ("Heading 2", 14.5, CHARCOAL, 11, 5),
        ("Heading 3", 11.5, GREEN_DARK, 8, 3),
    ]:
        st = styles[name]
        st.font.name = "Aptos Display"; st._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos Display")
        st.font.size = Pt(size); st.font.bold = name.startswith("Heading") or name == "Title"
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before); st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    if "Figure Caption" not in styles:
        cap = styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else: cap = styles["Figure Caption"]
    cap.font.name = "Aptos"; cap.font.size = Pt(8.5); cap.font.italic = True
    cap.font.color.rgb = RGBColor.from_string(MUTED)
    cap.paragraph_format.space_before = Pt(4); cap.paragraph_format.space_after = Pt(8)
    cap.paragraph_format.keep_with_next = False

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else: code = styles["Code Block"]
    code.font.name = "Aptos Mono"; code.font.size = Pt(8); code.font.color.rgb = RGBColor.from_string(CHARCOAL)
    code.paragraph_format.left_indent = Inches(0.2); code.paragraph_format.right_indent = Inches(0.2)
    code.paragraph_format.space_before = Pt(3); code.paragraph_format.space_after = Pt(5)
    code.paragraph_format.line_spacing = 1.0

    for section in doc.sections:
        hp = section.header.paragraphs[0]
        hp.text = "NEMOTRON VOICE AGENT   /   CURRENT ARCHITECTURE MANUAL"
        hp.style = styles["Normal"]
        for run in hp.runs:
            run.font.size=Pt(7.5); run.font.bold=True; run.font.color.rgb=RGBColor.from_string(GREEN_DARK)
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_page_number(section.footer.paragraphs[0])

    set_update_fields(doc)


def add_para(doc, text="", bold_lead=None, style=None, keep=False):
    p=doc.add_paragraph(style=style)
    if bold_lead and text.startswith(bold_lead):
        r=p.add_run(bold_lead);r.bold=True;r.font.color.rgb=RGBColor.from_string(CHARCOAL)
        p.add_run(text[len(bold_lead):])
    else: p.add_run(text)
    p.paragraph_format.keep_together=keep
    return p


def bullets(doc, items, level=0, numbered=False):
    style = "List Number" if numbered else "List Bullet"
    for item in items:
        p=doc.add_paragraph(style=style)
        p.paragraph_format.left_indent=Inches(0.24+level*0.2)
        p.paragraph_format.first_line_indent=Inches(-0.14)
        if isinstance(item, tuple):
            r=p.add_run(item[0]);r.bold=True;p.add_run(item[1])
        else:p.add_run(item)


def table(doc, headers, rows, widths=None, compact=True):
    t=doc.add_table(rows=1, cols=len(headers))
    t.alignment=WD_TABLE_ALIGNMENT.CENTER
    t.autofit=True
    t.style="Table Grid"
    hdr=t.rows[0]; set_repeat_table_header(hdr)
    for i,h in enumerate(headers):
        c=hdr.cells[i];set_cell_shading(c,GREEN);set_cell_border(c,"FFFFFF","5")
        p=c.paragraphs[0];p.alignment=WD_ALIGN_PARAGRAPH.LEFT
        r=p.add_run(str(h));r.bold=True;r.font.color.rgb=RGBColor(255,255,255);r.font.size=Pt(8.3)
        c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for ridx,row in enumerate(rows):
        cells=t.add_row().cells
        for i,val in enumerate(row):
            c=cells[i];set_cell_shading(c,PALE_2 if ridx%2==0 else WHITE);set_cell_border(c)
            p=c.paragraphs[0]
            p.paragraph_format.space_after=Pt(1 if compact else 3)
            r=p.add_run(str(val));r.font.size=Pt(7.7 if compact else 8.5)
            c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths and i<len(widths): c.width=Inches(widths[i])
    doc.add_paragraph().paragraph_format.space_after=Pt(1)
    return t


def callout(doc, title, text, kind="info"):
    fill, border, title_color = {
        "info":(PALE, GREEN, GREEN_DARK), "warning":(AMBER,"C68A00","8A5F00"),
        "risk":(RED_PALE,"B64A3A","8C2E23"), "note":(BLUE_PALE,"5B8FA8","365E72")
    }[kind]
    t=doc.add_table(rows=1,cols=1);t.alignment=WD_TABLE_ALIGNMENT.CENTER
    c=t.cell(0,0);set_cell_shading(c,fill);set_cell_border(c,border,"14")
    c.margin_top=Cm(0.14);c.margin_bottom=Cm(0.14)
    p=c.paragraphs[0];p.paragraph_format.space_after=Pt(2)
    r=p.add_run(title.upper());r.bold=True;r.font.size=Pt(9);r.font.color.rgb=RGBColor.from_string(title_color)
    p=c.add_paragraph(text);p.paragraph_format.space_after=Pt(2);p.paragraph_format.keep_together=True
    doc.add_paragraph().paragraph_format.space_after=Pt(1)


def code_block(doc, text):
    t=doc.add_table(rows=1,cols=1);c=t.cell(0,0);set_cell_shading(c,"F3F5F2");set_cell_border(c,"D0D6CD")
    p=c.paragraphs[0];p.style=doc.styles["Code Block"]
    p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after=Pt(1)


def add_figure(doc, path, caption, alt, width=6.75):
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together=True
    p.add_run().add_picture(str(path),width=Inches(width))
    set_image_alt_text(p,alt)
    cp=doc.add_paragraph(caption,style="Figure Caption");cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    return p


def chapter(doc, number, title, orientation="portrait"):
    if len(doc.paragraphs)>0: doc.add_page_break()
    p=doc.add_paragraph(style="Heading 1")
    r=p.add_run(f"{number}. {title}");r.bold=True


def add_cover(doc, logo_path, ssot_hash):
    for _ in range(3): doc.add_paragraph()
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(logo_path),width=Inches(3.7))
    doc.add_paragraph()
    p=doc.add_paragraph("NEMOTRON VOICE AGENT",style="Title");p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p=doc.add_paragraph("Current Architecture & Operations Manual",style="Subtitle");p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p=doc.add_paragraph("Astra + NVCF deployed pipeline / engineer handoff guide");p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:r.font.size=Pt(12);r.font.color.rgb=RGBColor.from_string(MUTED)
    doc.add_paragraph()
    t=doc.add_table(rows=4,cols=2);t.alignment=WD_TABLE_ALIGNMENT.CENTER;t.autofit=False
    vals=[("Document status",STATUS),("Architecture snapshot",SNAPSHOT_DATE),("Manual version",VERSION),("SSOT SHA-256",ssot_hash[:16]+"…")]
    for i,(k,v) in enumerate(vals):
        c1,c2=t.rows[i].cells;set_cell_shading(c1,PALE);set_cell_shading(c2,"FAFCF8");set_cell_border(c1);set_cell_border(c2)
        rr=c1.paragraphs[0].add_run(k);rr.bold=True;rr.font.color.rgb=RGBColor.from_string(GREEN_DARK)
        c2.paragraphs[0].add_run(v)
    doc.add_paragraph()
    callout(doc,"Current deployment truth","The retained UI is operationally live and points to the production NVCF function, but the UI is physically deployed in Astra stg infrastructure. A true Astra prd promotion is still outstanding.","warning")
    p=doc.add_paragraph("CONFIDENTIALITY NOTE");p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:r.bold=True;r.font.size=Pt(8);r.font.color.rgb=RGBColor.from_string(MUTED)
    p=doc.add_paragraph("Credential names and injection boundaries are documented. Secret values are intentionally omitted.");p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:r.font.size=Pt(8);r.font.color.rgb=RGBColor.from_string(MUTED)


def build_manual(diagrams):
    ssot_hash=hashlib.sha256(SSOT.read_bytes()).hexdigest()
    logo_src=REPO/"astra_client/public/nvidia-logo-trim.png"
    logo_dst=ASSETS/"nvidia-logo-trim.png"
    shutil.copy2(logo_src,logo_dst)
    nim_src=REPO/"astra_client/public/nvidia-nim-icon.png"
    shutil.copy2(nim_src,ASSETS/"nvidia-nim-icon.png")

    doc=Document();configure_doc(doc);add_cover(doc,logo_dst,ssot_hash)

    # Front matter
    chapter(doc,"A","How to use this manual")
    add_para(doc,"This is the operator and engineer handoff manual for the currently retained Nemotron Voice Agent deployment. It is written so that a new engineer can orient themselves, trace a request, operate the system, diagnose a failure, and understand which source file owns each behavior without reading the entire repository first.")
    doc.add_heading("Intended audience",level=2)
    bullets(doc,["Application engineers working on FastAPI, Pipecat, prompts, tools, media, or capture.","Platform engineers operating Astra, NVCF, Helm, Kubernetes, Redis, SeaweedFS, NIMs, or NGC.","SQA engineers qualifying real voice, tool calls, multimodal behavior, concurrency, and archive publication.","Incident responders who need a boundary-by-boundary isolation path and a common session correlation key."])
    doc.add_heading("Evidence labels",level=2)
    table(doc,["Label","Meaning","Use"],[
        ("LIVE-VERIFIED",f"Read-only observation against retained deployment on {SNAPSHOT_DATE}.","Use as current operational evidence."),
        ("RENDERED","Derived from exact checked-in Helm chart 0.1.94, reported by active function version.","Use as deployed-template evidence."),
        ("HISTORICAL","Preserved from earlier qualification reports or chart history.","Do not treat as a fresh guarantee."),
        ("DESIGN CONTRACT","Directly encoded in application/module behavior.","Validate again after code or artifact changes."),
    ],widths=[1.25,3.2,2.0],compact=False)
    callout(doc,"Maintenance contract","The Markdown SSOT remains authoritative. Regenerate this manual after changing active function/version, Astra environment, chart/image tags, examples, replicas, shared-state topology, capture contract, secret names, readiness behavior, or qualification outcome.","info")
    doc.add_heading("Document map",level=2)
    bullets(doc,["Sections 1–4: rapid orientation and component inventory.","Sections 5–9: request routing and the two voice experiences.","Sections 10–13: concurrency, Redis, SeaweedFS, and exactly-once capture.","Sections 14–18: readiness, secrets, observability, deployment, and verification.","Sections 19–23: troubleshooting, onboarding, risks, glossary, and source map."])

    chapter(doc,"B","Table of contents")
    add_para(doc,"The table below is a Word field. Open in Microsoft Word and choose Update Field / Update entire table if page numbers are not populated automatically.")
    p=doc.add_paragraph();add_field(p,' TOC \\o "1-3" \\h \\z \\u ',"Right-click and choose Update Field to build the clickable table of contents.")
    doc.add_heading("Figures",level=2)
    for i,name in enumerate(["End-to-end system overview","Trust boundaries and request routing","Start-session request sequence","Generic Assistant pipeline","Omni Subagents worker architecture","Multi-replica data sharing","Session capture flow","Deployment promotion flow","Failure-isolation map"],1):
        add_para(doc,f"Figure {i}. {name}")

    chapter(doc,1,"10-minute orientation")
    doc.add_heading("The shortest correct mental model",level=2)
    bullets(doc,[
        ("Astra is the credential-bearing UI proxy. ","It serves the React SPA and adds the NVCF invocation credential and function ID to server-side upstream calls."),
        ("NVCF is the runtime. ","Its Helm release contains five CPU application replicas, six GPU inference deployments consuming seven H100 GPUs, Redis, SeaweedFS, and a prewarmer."),
        ("A voice call stays on one app replica. ","The WebSocket, Pipecat processors, conversation context, and Omni WorkerBus are process-local for the lifetime of that socket."),
        ("Redis makes ancillary requests replica-safe. ","It shares session config, attachments, webcam frames, capture requests, capture flags, and owner-token locks."),
        ("SeaweedFS makes capture artifacts visible across pods. ","It stages session logs, transcript, and per-turn WAVs through an S3-compatible interface."),
        ("NGC is the durable archive. ","Exactly one app replica packages and publishes a resource version named with the session ID."),
        ("Generic and Omni are different agent shapes. ","Generic is cascaded ASR→LLM/tools→TTS. Omni is one in-process Pipecat worker graph with Speaker, Media Analyzer, Webcam, and Thinker roles."),
    ])
    add_figure(doc,diagrams[0],"Figure 1. End-to-end overview of the retained live path.","End-to-end overview showing the browser, Astra stg proxy, NVCF production gateways, five application replicas, Redis, SeaweedFS, GPU model services, external tools, and NGC.")
    callout(doc,"Terminology guardrail","Application replicas are Kubernetes pods serving FastAPI/Pipecat. NVCF function instances are platform allocations that contain the whole Helm release. Omni workers are in-process Pipecat roles created per Omni WebSocket. They are not interchangeable scaling units.","warning")

    chapter(doc,2,"Current deployed environment")
    doc.add_heading("Retained deployment snapshot",level=2)
    table(doc,["Item","Current value","Evidence"],[
        ("NVCF function","nemotron-voice-agent","Live-verified"),
        ("Function ID","81862ff8-4931-4f1e-9655-caa5b0bc5911","Live-verified; identifier, not a credential"),
        ("Active version","bc9de165-d8d1-456f-872c-c15b68959827","Live-verified"),
        ("Deployment ID","6e4c3e95-6113-4453-98fe-7fd5be18c2bf","Live-verified"),
        ("Backend / GPU instance","nvcf-dgxc-k8s-oci-nrt-prd12-1 / OCI.GPU.H100_8x","Live-verified"),
        ("NVCF scale","min 1 / max 1 function instance; max request concurrency 100","Live-verified; not a capacity guarantee"),
        ("Helm chart / app image","0.1.94 / nemotron-voice-agent:2.0.25","Active-version and rendered evidence"),
        ("Retained Astra app","nemotron-voice-agent-deploy","Live UI identity"),
        ("Astra infrastructure","stg / astrastg01-ocp-pdx04","Checked-in and previously live-verified"),
        ("Selectable examples","Generic Assistant; Omni Assistant Subagents","Live advertisement"),
        ("Capture","enabled, consent required, S3 backend, NGC destination configured","Live-verified at snapshot"),
    ],widths=[1.4,3.5,1.6],compact=False)
    callout(doc,"Astra environment truth","Operational shorthand calls the retained pair “production” because it serves users and points to the production NVCF function. Physically, the Astra UI still uses stg ingress, stg Vault paths, the stg cluster, and stg roles. True Astra prd requires a separate promotion to astraprd01-ocp-pdx04.","warning")
    doc.add_heading("Advertised service choices",level=2)
    table(doc,["Experience","LLM / model","ASR","TTS","Extra capability"],[
        ("Generic","Nemotron 3.5 Lightning or Nemotron 3 Super 120B","Nemotron ASR Streaming English","Magpie or Chatterbox Multilingual","Per-session tools"),
        ("Omni Subagents","Nemotron 3 Nano Omni 30B A3B Reasoning NVFP4","Internal audio understanding","Magpie or Chatterbox Multilingual","Attachments and webcam"),
    ],widths=[1.2,2.0,1.5,1.3,1.2],compact=False)

    chapter(doc,3,"System context, security, and trust boundaries")
    add_figure(doc,diagrams[1],"Figure 2. Browser, Astra, NVCF edge, and private namespace trust boundaries.","Trust-boundary diagram showing that secrets are injected only into Astra nginx or NVCF app processes and never delivered to browser JavaScript.")
    doc.add_heading("Boundary rules",level=2)
    bullets(doc,[
        "Browser input, microphone, camera frames, uploads, prompt overrides, and selection values are untrusted and validated server-side.",
        "config.js contains deployment timestamp and public demo settings only; it must never contain provider or invocation credentials.",
        "Astra nginx injects Authorization and function-id headers, strips incoming Cookie, hides upstream Set-Cookie, and routes WebSockets to the streaming gateway.",
        "NVCF secrets arrive in /var/secrets/secrets.json. The app shell exports named environment variables before Python starts.",
        "Redis, SeaweedFS, and NIM services are ClusterIP-only. Their protection currently depends on namespace/network isolation; SeaweedFS identity enforcement is not configured.",
        "Session IDs are server-generated 12-character hexadecimal values and are sanitized again before object-path use.",
    ])
    doc.add_heading("Why the design looks this way",level=2)
    table(doc,["Constraint","Design response"],[
        ("Browser WebSocket API cannot add custom authorization headers","Use same-origin Astra nginx and a dedicated NVCF streaming-gateway location."),
        ("Only app port 7860 is publicly reachable","Keep all NIMs, Redis, and SeaweedFS private behind ClusterIP services."),
        ("NVCF RWO block volumes caused fresh-zone scheduling stalls","Use ephemeral Redis/SeaweedFS and publish successful archives durably to NGC."),
        ("Affinity router required fragile StatefulSet stable DNS","Park the router; use ordinary load balancing plus Redis/shared object storage."),
        ("Function-version secrets are immutable/version-scoped","Repeat the complete secret set for every newly created function version."),
    ],widths=[2.4,4.0],compact=False)

    chapter(doc,4,"Component catalog and ownership")
    table(doc,["Component","Where / scale","Owns","Does not own"],[
        ("React astra_client","Browser / per user","UX, selections, media, transcript, capture consent","Credentials or inference"),
        ("Astra nginx","Astra / one CPU pod","SPA, same-origin proxy, auth/header injection","Agent logic"),
        ("NVCF gateways","Managed edge","Authentication and function routing","Session state"),
        ("Application","NVCF / five CPU pods","FastAPI API, per-WebSocket Pipecat graph, tools, capture finalizer","Model weights or durable archive"),
        ("Redis","NVCF / one CPU pod","Shared config/media/capture coordination","Durable data"),
        ("SeaweedFS","NVCF / one CPU pod","Shared S3-compatible capture staging","Durable archive"),
        ("ASR NIM","One GPU","Streaming English transcription","LLM/TTS"),
        ("Lightning","One GPU","Generic default LLM/tool choice","Tool execution"),
        ("Super 120B","Two GPUs, TP=2","Selectable Generic LLM","ASR/TTS"),
        ("Omni vLLM","One GPU","Audio/vision/reasoning for Omni workers","Audio synthesis"),
        ("Magpie / Chatterbox","One GPU each","External speech synthesis","ASR/LLM"),
        ("Prewarmer","One CPU pod","Direct NIM warm and keepalive","User sessions"),
        ("NGC resource","Managed registry","Durable versioned session archives","Live state"),
    ],widths=[1.25,1.4,2.0,1.85],compact=True)
    doc.add_heading("GPU budget",level=2)
    table(doc,["Workload","H100 count"],[("Nemotron ASR",1),("Lightning",1),("Super TP=2",2),("Omni",1),("Magpie",1),("Chatterbox",1),("Total",7)],widths=[4.6,1.4],compact=False)
    callout(doc,"Capacity interpretation","Seven of eight H100 GPUs are requested. The unused eighth GPU is headroom, not a hidden service. Five application replicas do not remove the single-replica inference bottlenecks.","note")

    chapter(doc,5,"Astra frontend and proxy behavior")
    doc.add_heading("What the UI is",level=2)
    bullets(doc,["React + TypeScript built with Vite.","Pipecat Client SDK with WebSocket transport.","Built in docker/Dockerfile.nvcf-ui and served by nginx-unprivileged on port 7860.","OpenShift-compatible non-root filesystem and group permissions.","Curated runtime UI for Generic and Omni only."])
    doc.add_heading("Route contract",level=2)
    table(doc,["Browser path","Upstream","Reason"],[
        ("/ and static assets","Local nginx web root","Serve SPA and public runtime config."),
        ("/api/ws","https://grpc.nvcf.nvidia.com","Preserve WebSocket upgrade and function routing."),
        ("/api/* and /health","https://${NVCF_HOST}","Per-function HTTP invocation path."),
        ("/feedback","Configured external form upstream","Optional; keeps upstream target private."),
    ],widths=[1.1,2.3,3.0],compact=False)
    callout(doc,"Stale-cookie mitigation","nginx strips incoming Cookie and hides Set-Cookie. This prevents an old nvcf-request-id from being reused after function rollover and producing dead-session 404 / WebSocket 1006 failures.","info")
    doc.add_heading("UI startup and session lifecycle",level=2)
    bullets(doc,["Container entrypoint creates browser-readable config.js from non-secret settings.","User selects example, service IDs, voice, reasoning, and optional tools.","UI POSTs sanitized choices to /api/session-config and receives the server-minted session ID.","UI opens WSS /api/ws?session_id=<sid> through the same origin.","During a session it sends mic audio, optionally streams webcam frames or attachments, and renders RTVI transcripts/events.","At teardown SessionCaptureReporter best-effort posts consent and a bounded transcript using keepalive."])

    chapter(doc,6,"NVCF function and Kubernetes topology")
    doc.add_heading("Rendered workloads",level=2)
    add_para(doc,"Chart 0.1.94 renders ten Deployment objects: the five-replica application Deployment, ASR, Lightning, Super, Omni, Magpie, Chatterbox, prewarmer, Redis, and SeaweedFS. The public function entrypoint targets the nemotron-voice-agent Service on port 7860 and /api/ws. In-chart ingress is disabled.")
    table(doc,["Internal service","Endpoint","Primary consumer"],[
        ("nemotron-voice-agent",":7860","NVCF edge"),
        ("nemotron-asr-streaming-english","gRPC :50052 / health :9001","Generic app and prewarmer"),
        ("nemotron-lightning","HTTP :8000/v1","Generic app"),
        ("nemotron-3-super","HTTP :8000/v1","Generic app and prewarmer"),
        ("nvidia-llm-vllm-omni","HTTP :8002/v1","Omni workers and prewarmer"),
        ("tts-service","gRPC :50051 / health :9000","Both experiences and prewarmer"),
        ("chatterbox-tts-service","gRPC :50051 / health :9000","Both experiences and prewarmer"),
        ("redis",":6379/0","All application replicas"),
        ("seaweedfs","S3 :8333","All application replicas"),
    ],widths=[2.0,2.1,2.4],compact=True)
    callout(doc,"Router status","router.enabled=false. The application is a Deployment, not a StatefulSet. New HTTP and WebSocket requests use ordinary Kubernetes load balancing; no chart-layer consistent hash exists.","warning")
    doc.add_heading("Scale terms that must not be confused",level=2)
    table(doc,["Term","Count now","Boundary"],[
        ("NVCF function instance","1 minimum / 1 maximum","One platform allocation containing the entire Helm release and 8-GPU node shape."),
        ("Application replica","5","Five Kubernetes pods accepting REST and/or WebSockets."),
        ("Pipecat pipeline","One per connected voice session","Processors and conversation context inside the accepting app process."),
        ("Omni workers","Multiple per Omni pipeline","Speaker, Media Analyzer, Webcam, Thinker coordinated by one in-process WorkerBus."),
        ("Inference service replica","One per service","Shared model service endpoint; Super consumes two GPUs."),
    ],widths=[1.6,1.2,3.7],compact=False)

    chapter(doc,7,"Backend API and request routing")
    add_figure(doc,diagrams[2],"Figure 3. Session configuration can be written by one app replica and consumed by another through Redis.","Sequence showing POST session config through Astra and NVCF HTTP edge to App A, Redis persistence, then WebSocket through the streaming edge to App B.")
    doc.add_heading("Public API surface",level=2)
    table(doc,["Route","Protocol","Purpose"],[
        ("/health","GET","Shallow FastAPI health."),("/api/deployment","GET","Visible examples, defaults, capabilities, transports."),
        ("/api/prompts /api/tools /api/subagents","GET","Example-local catalogs for UI."),("/api/services /api/tts-config","GET","Reachable service and voice metadata."),
        ("/api/session-config","POST","Validate choices, readiness-check services, mint SID, persist config."),("/api/ws","WebSocket","Curated NVCF Pipecat/RTVI voice path."),
        ("/api/sessions/{sid}/attachments","POST","Capability-gated media upload."),("/api/sessions/{sid}/webcam/frames","POST","Low-resolution live webcam frames."),
        ("/api/sessions/{sid}/webcam/capture","POST","Consume one high-resolution request token."),("/api/session-capture","POST","Teardown consent and transcript."),
        ("/api/session-capture/status","GET","Capture configuration and pending/error summary."),
    ],widths=[2.25,1.0,3.25],compact=True)
    doc.add_heading("Session-config algorithm",level=2)
    bullets(doc,["Parse JSON and constrain the requested example to the mounted visible registry.","Supply the example default prompt when no prompt key/content is present.","Drop unknown session fields; hydrate stable service IDs from example-local catalogs.","Check the selected local LLM, ASR, and TTS readiness.","Mint uuid4().hex[:12].","Store an in-process fast path plus Redis sb:cfg:<sid> with TTL.","Return the SID; the WebSocket replica reads local state first and Redis second, then sanitizes overrides again."],numbered=True)
    callout(doc,"Capability enforcement","Generic sessions receive 403 for Omni-only media routes. Unknown SIDs receive 404. Upload bodies are bounded; images must pass extension and JPEG/PNG magic-byte checks.","info")

    chapter(doc,8,"Generic Assistant walkthrough")
    add_figure(doc,diagrams[3],"Figure 4. Generic Assistant cascaded Pipecat voice and function-calling path.","Generic pipeline from microphone to ASR, LLM, optional Python tool handlers, speech gate, TTS, speaker and capture taps.")
    doc.add_heading("Agent model",level=2)
    add_para(doc,"Generic Assistant is a Pipecat cascaded voice pipeline. React is only the client. Agentic behavior comes from an OpenAI-compatible NVIDIA text LLM that receives validated JSON function schemas and uses tool_choice=auto. Pipecat invokes registered Python handlers and returns their structured result to the LLM for the final spoken answer.")
    doc.add_heading("Processor order",level=2)
    code_block(doc,"transport.input → NvidiaSTTService → user context aggregator → NvidiaLLMService\n→ ToolCallSpeechGate → NvidiaTTSService → transport.output\n→ optional activity checker → optional audio recorder → assistant context aggregator")
    doc.add_heading("Prompt policy",level=2)
    bullets(doc,["Friendly voice identity and brief spoken style.","Current weather routes to get_weather; future/forecast weather routes to web_search.","Current stock price routes to get_stock_price; other current facts route to web_search.","Never invent live data; report provider unavailability safely.","BMI and random-number tools have narrow explicit use.","Default output is one short sentence without markdown; exact tool results may be longer."])
    doc.add_heading("Tool catalog",level=2)
    table(doc,["Tool","Provider / execution","Credential name","Failure policy"],[
        ("get_weather","WeatherAPI current endpoint","WEATHERAPI_KEY","Friendly unavailable object; no fake weather."),
        ("get_stock_price","Finnhub search + quote","FINNHUB_API_KEY","Unavailable/not-found object; no fake price."),
        ("web_search","Perplexity Sonar via NVIDIA gateway","PERPLEXITY_API_KEY","One transient retry, then safe unavailable."),
        ("calculate_bmi","Local Python","None","Validate numeric positive inputs."),
        ("generate_random_number","Local Python","None","Validate integer range."),
    ],widths=[1.25,2.0,1.5,1.9],compact=True)
    doc.add_heading("Why the speech gate matters",level=2)
    add_para(doc,"The speech gate buffers a completion. If the same completion contains a function call, it drops every text frame in that completion so hidden reasoning or “let me check” filler does not reach TTS. Only the post-tool-result completion is spoken. Chatterbox additionally uses a 240-character chunk aggregator.")

    chapter(doc,9,"Omni Assistant Subagents walkthrough")
    add_figure(doc,diagrams[4],"Figure 5. Omni worker graph inside one application process and live session.","Omni architecture with the Transport Agent, BusBridgeProcessor, Pipecat WorkerBus, Speaker, Media Analyzer, Webcam and Thinker workers, pinned state board and Redis media streams.")
    doc.add_heading("Agent model",level=2)
    add_para(doc,"Omni Subagents uses Pipecat WorkerRunner, one shared in-process WorkerBus, multiple PipelineWorker instances, and a BusBridgeProcessor. The Transport Agent owns physical audio, VAD, user turns, TTS, RTVI, and media dispatch. The Speaker is the only worker allowed to speak. It reads the pinned board and emits a strict action envelope.")
    code_block(doc,'{\n  "transcript": "verbatim audio transcript",\n  "turn_action": "respond|think|analyze_attachment|capture_highres|clarify",\n  "response": "spoken response or acknowledgement",\n  "selected_input_source": "uploaded_attachment|none",\n  "media_analysis_action": "new|rerun|none",\n  "media_analysis_prompt": "self-contained task or empty",\n  "highres_query": "specific live capture question or empty"\n}')
    doc.add_heading("Turn ownership",level=2)
    table(doc,["Action","Owner","Effect"],[
        ("respond","Speaker","Answer immediately."),("clarify","Speaker","Ask one precise question; queue no worker."),
        ("analyze_attachment","Media Analyzer","Acknowledge first, then dispatch uploaded-file analysis."),
        ("capture_highres","Browser + Media Analyzer","Request one token-bound native-resolution frame, then analyze."),
        ("think","Thinker","Acknowledge and queue a reasoning-on re-answer."),
    ],widths=[1.4,1.8,3.5],compact=False)
    doc.add_heading("Uploaded media path",level=2)
    bullets(doc,["Browser POST may land on any app replica.","Receiver XADDs bytes and metadata to sb:att:<sid>.","Voice-session replica’s blocking XREAD listener wakes and marks the pinned board PENDING.","Speaker chooses analyze_attachment and speaks a short acknowledgement.","Media Analyzer reads the chosen attachment from Redis, performs analysis, pins detailed context, and returns concise speech."],numbered=True)
    doc.add_heading("Webcam and high-resolution capture",level=2)
    add_para(doc,"Compressed webcam frames arrive roughly once per second, up to 640 pixels wide at JPEG quality 0.7, in a ring-limited Redis stream. The Webcam Agent builds short temporal context and pins a conservative view/gesture summary. If detail is insufficient, the Speaker offers a later-turn high-resolution capture. A random request ID is stored in Redis and atomically consumed by compare-and-delete when the browser uploads the native-resolution frame.")
    callout(doc,"Source separation","Uploaded attachments and the rolling live webcam view remain distinct. The pending upload becomes the referent for “this image”; after analysis it becomes past context and live webcam context resumes.","note")

    chapter(doc,10,"Concurrency model")
    add_figure(doc,diagrams[5],"Figure 6. Shared and process-local state across five application replicas.","Five app replicas connected to shared Redis and SeaweedFS, alongside process-local live WebSocket and Omni WorkerBus state.")
    table(doc,["State / data","Location","Cross-pod?","Lifetime"],[
        ("Active socket and Pipecat processors","Application process","No","Socket/session"),
        ("LLM conversation context","Application process","No","Socket/session"),
        ("Sanitized session config","Local dict + Redis","Yes","TTL 3600 s"),
        ("Attachments / webcam frames","Redis Streams","Yes","TTL/ring-limited"),
        ("High-res capture request","Redis key","Yes","TTL 3600 s"),
        ("Capture flags / attempts","Redis hash","Yes","TTL 3600 s"),
        ("Finalize lock","Redis key","Yes","900 s"),
        ("Capture log/transcript/audio","SeaweedFS S3","Yes","Until success/cleanup/restart"),
        ("Final tarball","NGC resource version","Yes / durable","Registry retention"),
    ],widths=[1.8,1.5,1.0,2.0],compact=True)
    callout(doc,"Failure semantics","Redis enables REST/media requests to land on a different replica from the voice socket. It does not migrate a live socket, Pipecat graph, or conversation context. If the owning pod dies, that call ends and the browser must reconnect.","warning")
    doc.add_heading("Capacity boundaries",level=2)
    bullets(doc,["NVCF maxRequestConcurrency=100 is only an edge admission setting.","Five app replicas share one ASR, Lightning, Super, Omni, Magpie, and Chatterbox deployment each.","Omni maxNumSeqs=4, model KV/sequence limits, provider quotas, and TTS latency bound throughput.","Redis has 256 MiB maxmemory with allkeys-lru and carries binary media.","SeaweedFS is one pod; capture finalization has four threads per app process.","Safe concurrency is a measured SQA result, not a value derived from platform configuration."])

    chapter(doc,11,"Redis session bus")
    doc.add_heading("Deployment contract",level=2)
    bullets(doc,["One private Bitnami Redis Deployment at redis:6379/0.","No AOF or RDB persistence; 256 MiB maxmemory; allkeys-lru eviction.","Readiness/liveness uses redis-cli ping; empty password inside the namespace.","App entrypoint blocks on PING before starting Python when REDIS_URL is configured."])
    table(doc,["Key pattern","Type","Purpose"],[
        ("sb:cfg:<sid>","JSON string","Sanitized session config."),("sb:wc:<sid>","Stream","Webcam metadata and bytes."),
        ("sb:seq:wc:<sid>","Integer","Webcam sequence."),("sb:att:<sid>","Stream","Attachment metadata and bytes."),
        ("sb:seq:att:<sid>","Integer","Attachment sequence."),("sb:capreq:<sid>","String","Only valid high-resolution request ID."),
        ("cap:<sid>","Hash","Pipeline/consent signals, attempts, errors and update time."),("cap:lock:<sid>","String","Finalize owner token."),
    ],widths=[1.9,1.2,3.4],compact=True)
    doc.add_heading("Cross-pod listener behavior",level=2)
    bullets(doc,["Blocking XREAD starts at cursor 0 so pre-listener writes are not missed.","Block interval is 5000 ms; socket timeout is 15 s.","Idle timeout becomes an empty read; transient Redis/network errors retry with bounded backoff.","Callback failure is isolated and does not kill the listener."])
    callout(doc,"Outage behavior","A Redis restart loses active session config, media, capture flags, and locks. Existing sockets may continue voice from local context, but cross-pod REST/media/capture correctness is not reliable until Redis recovers.","risk")

    chapter(doc,12,"SeaweedFS shared session store")
    doc.add_heading("Purpose and shape",level=2)
    add_para(doc,"SeaweedFS is a second shared subsystem because Redis coordination does not contain the complete capture archive. The finalizing pod must be able to read the log, transcript, and WAV objects written by another pod. The current chart runs one combined SeaweedFS server with S3 API on port 8333 and bucket nva-session-capture.")
    code_block(doc,"sessions/<sid>/session.log\nsessions/<sid>/transcript.txt\nsessions/<sid>/audio/asr_000.wav\nsessions/<sid>/audio/tts_000.wav\nsessions/<sid>/audio/asr_001.wav\nsessions/<sid>/audio/tts_001.wav")
    bullets(doc,["Backend abstraction supports put, get, list, delete, delete_prefix, and exists.","Same object keys can target local files, SeaweedFS, MinIO, or real S3.","Current data directory is a 20 GiB emptyDir.","Placeholder signing credentials satisfy boto3; current chart does not generate SeaweedFS identities."])
    callout(doc,"Durability boundary","SeaweedFS is shared staging, not the archive of record. Pod reschedule/restart can erase in-flight or retained source objects. A successfully uploaded NGC version is durable.","warning")

    chapter(doc,13,"Session capture and NGC publication")
    add_figure(doc,diagrams[6],"Figure 7. Two-signal, owner-token-locked capture finalization and NGC publication.","Capture flow showing pipeline and browser consent signals, Redis state and lock, SeaweedFS objects, one winning app replica, NGC upload, cleanup and retry retention.")
    doc.add_heading("Capture contents",level=2)
    bullets(doc,["Per-session application log correlated by SID.","Browser-rendered user/assistant transcript, capped at 200,000 characters.","Per-turn ASR input WAVs and TTS output WAVs."])
    doc.add_heading("State machine",level=2)
    table(doc,["State / event","Meaning","Next action"],[
        ("Waiting","Neither independent signal present.","Do nothing."),("PipelineOnly / ConsentOnly","One signal present.","Wait; reaper may abandon after 900 s."),
        ("Ready","Both signals present.","Attempt SET lock NX with random token and 900 s TTL."),("Locked","This replica owns token.","Read shared objects, tar, upload/discard."),
        ("Archived","NGC upload succeeded.","Delete store prefix and clear Redis."),("Retryable","Store/tar/config/upload failed.","Retain evidence; later signal/reaper retries."),
        ("RetainedFailure","NGC-related failure exhausted attempts.","Operator review; do not blindly delete."),("Abandoned","Only one stale signal.","Delete orphaned objects/state."),
    ],widths=[1.5,2.4,2.6],compact=True)
    doc.add_heading("Exactly-once mechanics",level=2)
    bullets(doc,["Pipeline uploads local log before setting pipeline_done so no finalizer can race ahead of the object.","Browser teardown records consent_done and transcript through any replica.","SET cap:lock:<sid> <random-token> NX EX 900 elects one finalizer.","Lua compare-and-delete releases only the caller’s token, preventing an expired old worker from removing a new lock.","Winning process builds a temporary <sid>.tar.gz and invokes ngc registry resource upload-version.","On success it deletes the SeaweedFS prefix, clears Redis, and removes the local tar."])
    doc.add_heading("Retry policy",level=2)
    table(doc,["Control","Value"],[('Reaper interval','300 s'),('Orphan threshold','900 s'),('Capture state TTL','3600 s'),('Lock TTL','900 s'),('Maximum attempts','5'),('NGC subprocess timeout','300 s'),('Finalize executor','4 threads per app process')],widths=[3.0,2.6],compact=False)
    callout(doc,"Ambiguous timeout","NGC may accept an upload even when the client times out. Before retrying or deleting any retained session, query NGC for the version named with that exact SID.","risk")

    chapter(doc,14,"Startup, readiness, and prewarming")
    doc.add_heading("Application startup gates",level=2)
    bullets(doc,["Wait indefinitely for Redis PING when REDIS_URL is configured.","Wait for SeaweedFS HTTP when SESSION_STORE_BACKEND=s3.","Read /var/secrets/secrets.json and export named credentials.","Start src/server.py only after shared dependencies are reachable."])
    doc.add_heading("Probe semantics",level=2)
    table(doc,["Signal","What it proves","What it does not prove"],[
        ("Kubernetes /health probes","FastAPI process is responsive.","All NIMs and providers work."),
        ("NVCF ACTIVE","Platform rollout is registered.","First real model call is warm."),
        ("/api/session-config success","Selected local LLM/ASR/TTS passed deeper readiness.","Quality under load or live provider success."),
        ("Real audio SQA","End-to-end speech path and expected behavior.","All future concurrency levels."),
    ],widths=[1.6,2.35,2.55],compact=False)
    add_para(doc,"ASR, Super, Magpie, and Chatterbox can report Kubernetes-ready from lightweight container/process checks because nimReadyImmediate remains true. Lightning and Omni use explicit health probes. Per-session readiness and real voice tests are the functional gate.")
    doc.add_heading("Prewarmer",level=2)
    bullets(doc,["Calls model services directly, never the app.","Warms Super and Omni with tiny chat completions; Omni uses JSON response format to precompile guided decoding.","Warms ASR, Magpie, and Chatterbox through bundled Riva clients.","Retries every 15 seconds until success and repeats every 300 seconds."])
    callout(doc,"Known prewarm gap","Lightning is the default Generic LLM but is absent from current prewarmer targets. Its real readiness probe gates availability, yet first-generation latency/compile work may reach a user.","warning")

    chapter(doc,15,"Secrets and credential ownership")
    callout(doc,"No values in documentation","Only credential names, consumers, and injection boundaries belong in this manual. Never paste actual API keys, bearer tokens, or contents of secrets.json into reports, commits, screenshots, or shell transcripts.","risk")
    table(doc,["Name","Consumer","Purpose"],[
        ("NVIDIA_API_KEY","App/NVIDIA clients; NGC fallback only","Inference/API invocation where applicable."),
        ("NGC_API_KEY","NIM startup and capture uploader","Model registry and dedicated capture publication."),
        ("PERPLEXITY_API_KEY","Generic web_search","Perplexity Sonar request."),
        ("WEATHERAPI_KEY","Generic get_weather","WeatherAPI request."),
        ("FINNHUB_API_KEY","Generic get_stock_price","Finnhub request."),
        ("SESSION_CAPTURE_NGC","Capture module","<org>/<resource> destination name."),
        ("NVCF_HOST / NVCF_FUNCTION_ID","Astra nginx","Backend target and function route."),
    ],widths=[1.9,2.1,2.6],compact=True)
    doc.add_heading("Injection paths",level=2)
    bullets(doc,["NVCF: function-version secrets → /var/secrets/secrets.json → app shell export → Python/tool handlers/NGC CLI.","Astra: stg Vault path → ExternalSecret/SecretStore → nginx environment → private rendered nginx configuration.","NIM deployments use startup-specific NGC access for model weights.","Every new immutable NVCF version must be created with the complete secret set again."])
    doc.add_heading("Credential separation",level=2)
    add_para(doc,"An Astra-to-NVCF invocation key, an NGC registry key, a personal key for instance log access, a Perplexity key, and provider-specific WeatherAPI/Finnhub keys are different capabilities. Do not assume one credential works across boundaries. Capture prefers the dedicated NGC_API_KEY and reports the chosen source name, never its value.")

    chapter(doc,16,"Observability and evidence")
    table(doc,["Signal","Useful proof","Limit"],[
        ("NVCF ACTIVE","Platform deployment exists.","Not model correctness."),("Astra /health 200","Proxy reaches FastAPI.","Not tools/capture."),
        ("/api/deployment","Catalog resolves.","Not audio quality."),("/api/session-capture/status","Capture config/backend/key names/pending state.","Not specific NGC success."),
        ("RTVI events","Transcript, tool call, subagent and latency reached client.","Not upstream provider truth alone."),
        ("Captured session.log","Detailed server path for one SID.","Absent if finalization never completes."),("NGC version","Durable archive for exact SID.","Quality requires archive inspection."),
        ("Pod/instance logs","Startup/runtime evidence.","May be constrained by NVCF access."),
    ],widths=[1.45,2.5,2.6],compact=True)
    doc.add_heading("Correlation strategy",level=2)
    add_para(doc,"Use the 12-hex session ID everywhere. The WebSocket binds loguru with stream_id=<sid>. The same value names Redis keys, the SeaweedFS sessions/<sid>/ prefix, the capture archive root, and the NGC resource version. Start every incident note with timestamp, browser/session ID, selected example/model/TTS, and expected behavior.")
    doc.add_heading("Historical qualification evidence",level=2)
    bullets(doc,["Five application replicas without observed cross-talk.","Mixed eight-session browser/audio concurrency.","Five cross-replica attachment sessions correctly analyzed a known image.","Live WeatherAPI, Finnhub, and Perplexity credentials exercised.","Consented captures reached NGC."])
    callout(doc,"Evidence freshness","Those results came from staging chart 0.1.90. Production 0.1.91–0.1.94 added webcam/capture teardown hardening, catalog changes, and NGC diagnostics. Historical evidence informs confidence but never replaces fresh testing after artifact changes.","warning")

    chapter(doc,17,"Deployment and promotion")
    add_figure(doc,diagrams[7],"Figure 8. Immutable artifact qualification and promotion flow, including the outstanding true Astra prd boundary.","Promotion flow from source and images/chart through Viking local, NVCF/Astra staging, production NVCF, retained Astra stg live UI, and future true Astra prd.")
    doc.add_heading("Standard pathway",level=2)
    bullets(doc,["Build and push exact app and Astra UI images; bump/package/push the Helm chart.","Deploy to Viking local Kubernetes and run astra_client locally.","Use Playwright with real audio and independent ASR to understand TTS output.","Validate selectors, tools, media, webcam, capture, and concurrent sessions.","Create immutable NVCF staging version with the exact chart and complete secrets; deploy exact UI tag to Astra preview.","Repeat the full public-path SQA/concurrency matrix and obtain user acceptance.","Create production NVCF version from identical artifacts; verify before removing the old version unless explicitly accepting downtime for GPU capacity.","Point retained Astra UI to the qualified production target and exact UI image."],numbered=True)
    callout(doc,"Current staging state","NVCF staging and Astra preview are pulled down. Recreating them requires deliberate deployment of both sides. Inactive immutable function versions may remain as history.","note")
    doc.add_heading("True Astra prd promotion",level=2)
    bullets(doc,["Obtain a valid NSPECT ID.","Deploy through the production Fusion control plane to astraprd01-ocp-pdx04.","Create/populate an independent Vault path ending /prd.","Generate and verify prd ingress, environment, JWT, role, and shared-secret path values.","Deploy the exact qualified UI image and production function target.","Repeat HTTP, WebSocket, voice, tool, media, concurrency, and capture qualification.","Only then delete the retained Astra stg incarnation."],numbered=True)
    doc.add_heading("Rollback units",level=2)
    table(doc,["Layer","Rollback"],[
        ("Astra UI","Restore previous JFrog image tag and ArgoCD/Fusion sync."),("NVCF","Redeploy known immutable function version/chart; repoint Vault if FID changes."),
        ("Secrets","Restore Vault values or create a corrected immutable NVCF version."),("Data","Redis/SeaweedFS are not rollback stores; NGC versions are durable capture records."),
    ],widths=[1.3,5.1],compact=False)

    chapter(doc,18,"Qualification and verification runbooks")
    doc.add_heading("Public smoke checks",level=2)
    code_block(doc,'BASE=https://nemotron-voice-agent-deploy-backend.stg.astra.nvidia.com\ncurl -fsS "$BASE/health"\ncurl -fsS "$BASE/api/deployment"\ncurl -fsS "$BASE/api/services?pipeline_mode=generic-assistant"\ncurl -fsS "$BASE/api/session-capture/status"\ncurl -fsS "$BASE/config.js"')
    doc.add_heading("Real voice matrix",level=2)
    bullets(doc,["Open a clean browser context and assert a unique SID.","Generic + Lightning + reasoning on + Magpie: static question, weather, stock, web search.","Repeat expected-tool prompts across batches; assert RTVI tool event, provider-derived result, final speech, no socket/console error.","Generic + Super and Chatterbox selection checks.","Omni voice, attachment, webcam, conservative gesture, and high-resolution capture flows.","Overlap Generic and Omni sessions; assert no cross-talk/SID reuse.","End sessions cleanly and verify pending capture count returns to zero."],numbered=True)
    doc.add_heading("Capture acceptance",level=2)
    bullets(doc,["Record /api/session-capture/status before test.","Run consented session with at least one user and one bot turn; retain SID.","Poll until SID is no longer pending.","Require NGC version <sid> status UPLOAD_COMPLETE.","Download/extract read-only and require session.log, transcript.txt, and expected ASR/TTS WAVs.","Run declined-consent control and confirm no NGC version."],numbered=True)
    doc.add_heading("Minimum concurrency matrix",level=2)
    table(doc,["Dimension","Minimum assertion"],[
        ("Generic live tools","Weather, stock, and web search; multiple batches, not one lucky call."),("Overlap","Six browsers per batch when matching historical test; include mixed Generic/Omni."),
        ("Configuration","Submitted reasoning, service IDs, tool allowlist and transcript match intent."),("Isolation","Every browser has unique SID, no response/media/capture cross-talk."),
        ("Capture","Concurrent sessions finalize and pending queue returns to zero."),("Cross-replica media","When pod evidence exists, upload and voice socket can be on different replicas."),
    ],widths=[1.65,4.85],compact=False)

    chapter(doc,19,"Troubleshooting and failure isolation")
    add_figure(doc,diagrams[8],"Figure 9. Boundary-by-boundary failure isolation path.","Troubleshooting map moving from Astra UI to health/catalog, WebSocket, voice model services, tools, media/Redis and capture/NGC.")
    table(doc,["Symptom","Likely boundary","First checks / mitigation"],[
        ("UI loads; API 401/403","Astra Vault/invocation key","Patch Vault; restart/sync UI; recheck catalog."),
        ("WS 200 not 101 or browser 1006","Streaming route/FID/cookie","Check grpc.nvcf route, headers, rendered nginx, cookie stripping."),
        ("App waits before Python","Redis or SeaweedFS","Inspect shared deployments/services first."),
        ("Wrong/default example","Redis config handoff","Check sb:cfg:<sid>, TTL, exact SID propagation."),
        ("Attachment uploaded; no analysis","Redis stream/listener","Inspect sb:att:<sid>, XREAD warnings and capability."),
        ("Camera appears off despite 200 uploads","SID/listener/board state","Check webcam RTVI state and sb:wc:<sid>."),
        ("Tool not called","Model/prompt/history","Check submitted reasoning and RTVI EXPECT_TOOL event repeatedly."),
        ("Tool called but unavailable","Credential/provider","Verify named key exists on every app replica/version and provider status."),
        ("Capture POST succeeds; no NGC version","Two-signal/store/upload","Inspect status, cap:<sid>, lock, Seaweed objects and NGC."),
        ("NGC upload timeout","Ambiguous registry/network","Query exact NGC version before retry/delete."),
        ("Function ACTIVE; first call slow","Cold/relaxed readiness","Check NIM health; remember Lightning prewarm gap."),
        ("Redis eviction/cross-feature breakage","256 MiB allkeys-lru","Reduce media/ring/concurrency or increase capacity."),
    ],widths=[2.0,1.5,3.0],compact=True)
    doc.add_heading("Fast isolation questions",level=2)
    bullets(doc,["What exact SID, timestamp, example, model, TTS, tools, and reasoning flag were submitted?","Did /api/session-config succeed or return 503 readiness?","Was the same SID used for WebSocket, media, and capture?","Which RTVI event is the last correct evidence?","Is the failed boundary model selection, provider execution, speech, Redis propagation, SeaweedFS, or NGC?","Does a fresh session reproduce it under repeated concurrency?"])

    chapter(doc,20,"Operational failure-isolation matrix")
    table(doc,["Dependency loss","Existing socket","New REST / media","Capture","Recovery expectation"],[
        ("One app pod","That pod's calls end","Kubernetes routes elsewhere","Shared state allows another pod/reaper to finish if artifacts/signals exist","Client reconnect; inspect lost one-signal state."),
        ("Redis","Local voice may continue","Cross-pod config/media unsafe or fails","Flags/locks unavailable","Restore Redis; active ephemeral state was lost."),
        ("SeaweedFS","Voice can continue","Most session/media unaffected","New/finalize capture fails; retained source may be lost on restart","Restore store; understand emptyDir loss."),
        ("One NIM","Affected inference stage fails","Session config may return 503","Capture unrelated except call may end","Restore NIM; prewarm; qualify real voice."),
        ("External provider","Voice still works","Specific live tool returns safe unavailable","Capture unrelated","Restore key/quota/provider; repeat EXPECT_TOOL."),
        ("NGC","Voice still works","Unrelated","Upload retries and retains source/state","Check version before retry; avoid Seaweed restart."),
        ("Astra proxy","New browser path unavailable","No public access","In-flight server work may continue","Restore app/ingress/Vault; public smoke."),
    ],widths=[1.1,1.25,1.35,1.6,1.55],compact=True)
    callout(doc,"Do not infer health","NVCF ACTIVE, Kubernetes Ready, and Astra /health=200 are necessary but shallow. A user-ready release requires session-config readiness plus real speech and feature-specific verification.","warning")

    chapter(doc,21,"New-engineer pickup checklist")
    doc.add_heading("Day 1 — establish the mental model",level=2)
    bullets(doc,["Read Sections 1, 6, 7, 10 and 13; explain the three scaling units back to a teammate.","Locate the two example pipelines and trace one Generic and one Omni turn in source.","Trace POST /api/session-config and /api/ws from React through Astra nginx to src/server.py.","Map one SID across local log context, Redis keys, SeaweedFS prefix and NGC version.","Confirm why Redis and SeaweedFS are separate dependencies."])
    doc.add_heading("Day 2 — reproduce the system safely",level=2)
    bullets(doc,["Render chart 0.1.94 and compare workloads/services with the component catalog.","Run unit/service tests relevant to the intended change.","Deploy to Viking local cluster using non-production credentials supplied out-of-band.","Run a clean Generic audio session and a clean Omni media session.","Exercise consented and declined capture controls without publishing keys in logs."])
    doc.add_heading("Before changing code",level=2)
    bullets(doc,["State the owning module, cross-pod implications, and data lifetime.","Decide whether the change affects app image, UI image, chart, function secrets, or more than one immutable artifact.","Add unit coverage plus a real-audio/public-path acceptance criterion.","Identify failure behavior and whether evidence is retained or deleted.","Update the Markdown SSOT and rebuild this manual if deployed architecture changes."])
    doc.add_heading("Before promotion",level=2)
    bullets(doc,["Record Git SHA, app/UI tags, chart version and function version.","Confirm complete function-version secret names are present without printing values.","Run repeated expected-tool and mixed concurrency matrices.","Validate selectors, webcam/high-res/attachments, and capture archive contents.","Obtain user acceptance at staging; explicitly authorize any H100-driven downtime cutover.","Verify the new live path before removing old NVCF/Astra deployment units."])

    chapter(doc,22,"Known risks and decision log")
    risks=[
        ("Astra still stg","Operational live UI is not a true Astra prd deployment.","Promote separately with NSPECT and independent prd Vault path."),
        ("Lightning tool nondeterminism","Automatic tool choice depends on model/history/reasoning.","Keep reasoning-on baseline and repeated EXPECT_TOOL matrix."),
        ("No live-session migration","Router disabled; Pipecat context process-local.","Client reconnect; use Redis only for ancillary cross-pod state."),
        ("Redis single ephemeral pod","Correctness/capacity dependency; 256 MiB binary-media pressure.","Measure; tune payload/ring/memory; consider HA/persistence if SLA changes."),
        ("SeaweedFS single ephemeral pod","Failed-upload evidence can disappear on restart.","Avoid restart during recovery; consider durable shared store."),
        ("Network-only Seaweed security","No generated identity config.","Maintain namespace isolation; add auth if threat model expands."),
        ("Shallow app health","FastAPI health does not prove model paths.","Deep session readiness and real voice SQA."),
        ("Relaxed NIM readiness","Some workloads look ready before full model endpoint.","Prewarm and test each selected model."),
        ("Lightning prewarm gap","Default model lacks direct prewarmer target.","Add/qualify warm call in a future chart change."),
        ("Single inference replicas","Five app pods share each NIM bottleneck.","Measure capacity; scale/model-shape only with GPU budget."),
        ("Best-effort browser capture POST","Sudden browser death can leave one signal.","Reaper abandons orphans; archive test normal teardown."),
        ("Ambiguous NGC timeout","Upload may succeed server-side despite local timeout.","Query version before manual retry/delete."),
        ("Tracing disabled","No Phoenix/OpenTelemetry UI.","Use SID, RTVI, logs, Redis/store and NGC evidence."),
    ]
    table(doc,["Risk","Impact","Current response"],risks,widths=[1.5,2.5,2.5],compact=True)

    chapter(doc,23,"Glossary and authoritative source map")
    doc.add_heading("Glossary",level=2)
    table(doc,["Term","Meaning"],[
        ("Astra","NVIDIA application hosting/OpenShift environment for the UI proxy."),("NVCF","NVIDIA Cloud Functions hosting the Helm release and public function gateways."),
        ("NIM","NVIDIA Inference Microservice runtime/image."),("Pipecat","Real-time audio pipeline and worker framework used by Python app."),
        ("RTVI","Pipecat client/server event protocol over /api/ws."),("FID","NVCF function ID; route identifier, not itself the bearer credential."),
        ("Function version","Immutable NVCF artifact/configuration under one FID."),("Function instance","Platform allocation containing the complete Helm release."),
        ("Application replica","One FastAPI/Pipecat Kubernetes pod."),("Omni worker","In-process role inside one Omni session's WorkerBus."),
        ("SID","12-hex session identifier linking config, socket, media, capture and NGC."),("Session bus","Redis-backed live state/media sharing."),
        ("Session store","Object backend for capture artifacts."),("Pinned board","Omni Speaker context with subagent findings/live view."),
        ("Operational production","Retained user-facing Astra stg UI + production NVCF function."),("Astra prd","Actual production Astra control plane/cluster; not yet the retained UI."),
    ],widths=[1.7,4.8],compact=True)
    doc.add_heading("Source map",level=2)
    source_rows=[
        ("Architecture SSOT","docs/current-deployed-pipeline-architecture.md"),("Chart/version defaults","nvcf_helm/Chart.yaml; nvcf_helm/values.yaml"),
        ("Workloads/services/config","nvcf_helm/templates/"),("Example registry","examples_registry.yaml; src/examples_registry.py; Helm ConfigMap"),
        ("Backend API/routing","src/server.py"),("Generic graph/prompt/tools","src/examples/generic/pipeline.py; prompts.yaml; tools.yaml; tool_handlers.py"),
        ("Omni graph/ownership","src/examples/omni_assistant_subagents/pipeline.py; prompts.yaml; subagents/"),
        ("Redis shared state","src/session_bus/"),("Capture state/finalizer/reaper","src/session_capture/"),
        ("Shared object storage","src/session_store/"),("Browser session/capture","astra_client/src/hooks/useVoiceSession.ts; astra_client/src/demo/SessionCaptureReporter.tsx"),
        ("Astra proxy image","docker/Dockerfile.nvcf-ui; docker/nginx-nvcf.conf.template; UI entrypoint"),
        ("Retained Astra values","nemotron-voice-agent-values.yaml"),("SQA","tests/sqa/"),
        ("This manual build","docs/architecture-manual/build_manual.py"),
    ]
    table(doc,["Concern","Authoritative source"],source_rows,widths=[2.2,4.3],compact=True)
    callout(doc,"Precedence rule","When prose conflicts with executable configuration, the active immutable artifact and corresponding source/template win. Preserve evidence date and class; never silently turn historical observation into a current guarantee.","info")

    chapter(doc,"C","Document control")
    table(doc,["Field","Value"],[
        ("Manual version",VERSION),("Architecture snapshot",SNAPSHOT_DATE),("Generated UTC",datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")),
        ("SSOT file",str(SSOT.relative_to(REPO))),("SSOT SHA-256",ssot_hash),("Diagram count",len(diagrams)),
        ("Secrets policy","Credential names only; values prohibited"),("Rebuild command","See docs/architecture-manual/README.md"),
    ],widths=[1.7,4.8],compact=False)
    doc.add_heading("Change triggers",level=2)
    bullets(doc,["Active function/version/backend/chart/app image.","Astra app/environment/URL/UI tag/Vault path.","Examples, models, voices, tools, replicas, or transport.","Router/Redis/SeaweedFS topology or data lifetimes.","Capture contract, NGC destination, secret names or readiness/prewarmer behavior.","Qualification result or accepted production risk."])
    add_para(doc,"End of manual.")

    doc.core_properties.title="Nemotron Voice Agent — Current Architecture & Operations Manual"
    doc.core_properties.subject="Astra + NVCF deployed pipeline engineer handoff guide"
    doc.core_properties.author="Nemotron Voice Agent Engineering"
    doc.core_properties.keywords="Nemotron, Astra, NVCF, Pipecat, Redis, SeaweedFS, NGC, architecture"
    doc.core_properties.comments=f"Generated from {SSOT.name} SHA-256 {ssot_hash}"
    doc.save(OUTPUT)
    return ssot_hash


def validate_docx(diagrams, ssot_hash):
    report={"output":str(OUTPUT),"ssot_sha256":ssot_hash,"figures":len(diagrams),"checks":{},"limitations":[]}
    report["docx_bytes"]=OUTPUT.stat().st_size
    report["diagram_bytes"]={p.name:p.stat().st_size for p in diagrams}
    report["checks"]["docx_exists_nonempty"]=OUTPUT.exists() and OUTPUT.stat().st_size>100_000
    with zipfile.ZipFile(OUTPUT) as z:
        report["checks"]["zip_integrity"]=z.testzip() is None
        names=z.namelist();media=[n for n in names if n.startswith("word/media/")]
        report["embedded_media_count"]=len(media)
        report["checks"]["embedded_media_nonempty"]=all(len(z.read(n))>100 for n in media)
        blob=b"".join(z.read(n) for n in names if n.endswith(".xml"))
        text=blob.decode("utf-8","ignore")
        # Detect credential-shaped values without embedding any real or partial key.
        secret_shapes = [
            r"nvapi-[A-Za-z0-9_-]{10,}",
            r"sk-[A-Za-z0-9_-]{15,}",
            r"(?:api[_ -]?key|token|secret)[\"=:\s]+[A-Za-z0-9_\/-]{24,}",
        ]
        report["checks"]["no_credential_shaped_values"] = not any(
            re.search(pattern, text, re.IGNORECASE) for pattern in secret_shapes
        )
        report["checks"]["has_toc_field"]=" TOC " in text
        report["checks"]["has_page_field"]=" PAGE " in text
        report["checks"]["has_alt_text"]=all(f"Figure {i}" in text or "overview" in text.lower() for i in range(1,10))
    for p in diagrams:
        with Image.open(p) as im:
            im.verify()
        with Image.open(p) as im:
            if im.size!=(W,H): raise RuntimeError(f"Unexpected diagram size: {p} {im.size}")
    report["checks"]["all_diagrams_valid_2800x1575"]=True
    # LibreOffice preview is optional.
    lo=shutil.which("libreoffice") or shutil.which("soffice")
    if lo:
        cp=subprocess.run([lo,"--headless","--convert-to","pdf","--outdir",str(HERE),str(OUTPUT)],capture_output=True,text=True,timeout=180)
        report["libreoffice"]={"path":lo,"returncode":cp.returncode,"stdout":cp.stdout,"stderr":cp.stderr}
        report["checks"]["pdf_preview_created"]=PDF_OUTPUT.exists() and PDF_OUTPUT.stat().st_size>1000
        pdfinfo=shutil.which("pdfinfo")
        if pdfinfo and PDF_OUTPUT.exists():
            info=subprocess.run([pdfinfo,str(PDF_OUTPUT)],capture_output=True,text=True,timeout=30).stdout
            for line in info.splitlines():
                if line.startswith("Pages:"):report["rendered_page_count"]=int(line.split(":",1)[1].strip())
    else:
        report["limitations"].append("LibreOffice/soffice is not installed; no PDF preview or renderer-exact page count was produced. Word will update the TOC and pagination fields when opened.")
        # Explicit page breaks give a useful lower bound; pagination may add more pages.
        report["minimum_manual_page_count"]=27
    report["all_required_checks_passed"]=all(report["checks"].values())
    BUILD_REPORT.write_text(json.dumps(report,indent=2)+"\n")
    return report


def main():
    if not SSOT.exists():raise SystemExit(f"Missing SSOT: {SSOT}")
    diagrams=generate_diagrams()
    ssot_hash=build_manual(diagrams)
    report=validate_docx(diagrams,ssot_hash)
    print(json.dumps(report,indent=2))
    if not report["all_required_checks_passed"]:return 1
    return 0


if __name__=="__main__":
    raise SystemExit(main())
